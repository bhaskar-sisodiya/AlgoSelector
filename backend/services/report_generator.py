import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
)
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from services.visualizer import Visualizer


# ────────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ────────────────────────────────────────────────────────────────────────────────

def _safe(value, decimals: int = 4) -> str:
    """Return a clean string representation of any value."""
    if value is None:
        return "N/A"
    if isinstance(value, float):
        return str(round(value, decimals))
    return str(value)


def _chart_to_tmp(buf: io.BytesIO) -> str | None:
    """Write a chart BytesIO to a temp PNG file; return path or None."""
    import tempfile
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(buf.getvalue())
            return tmp.name
    except Exception:
        return None


# ────────────────────────────────────────────────────────────────────────────────
# ReportGenerator
# ────────────────────────────────────────────────────────────────────────────────

class ReportGenerator:

    # ── DOCX ────────────────────────────────────────────────────────────────────

    @staticmethod
    def generate_docx(data: dict) -> io.BytesIO:
        doc = Document()
        automl = data.get("automl_results", {})

        # Title
        title = doc.add_heading(
            f"AutoML Report: {data.get('dataset_name', 'Dataset')}", 0
        )
        title.runs[0].font.color.rgb = RGBColor(30, 136, 229)

        # 1. Dataset Overview
        doc.add_heading("1. Dataset Overview", level=1)
        for label, value in [
            ("Rows", data.get("rows")),
            ("Columns", data.get("columns")),
            ("Task Type", data.get("task_type", "N/A")),
            ("Target Column", data.get("target_column", "N/A")),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(_safe(value))

        # 2. Meta-Insights
        doc.add_heading("2. Meta-Insights", level=1)
        meta = data.get("meta_features", {})
        if meta:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            tbl.rows[0].cells[0].text = "Metric"
            tbl.rows[0].cells[1].text = "Value"
            for name, val in [
                ("Dimensionality Ratio", meta.get("dimensionality_ratio")),
                ("Class Imbalance Ratio", meta.get("class_imbalance_ratio")),
                ("Avg Feature Correlation", meta.get("avg_feature_correlation")),
                ("Skewness", meta.get("skewness")),
            ]:
                row = tbl.add_row().cells
                row[0].text = name
                row[1].text = _safe(val)
        else:
            doc.add_paragraph("No meta-features available.")

        # 3. Algorithm Leaderboard
        doc.add_heading("3. Algorithm Leaderboard", level=1)
        models = automl.get("algorithms", [])
        if models:
            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Table Grid"
            for i, h in enumerate(["Rank", "Algorithm", "Accuracy", "F1 Score", "Time (s)"]):
                tbl.rows[0].cells[i].text = h
            for i, m in enumerate(models, 1):
                row = tbl.add_row().cells
                row[0].text = str(i)
                row[1].text = m.get("name", "Unknown")
                row[2].text = f"{m.get('accuracy', 0)}%"
                row[3].text = f"{m.get('f1_score', 'N/A')}%"
                row[4].text = _safe(m.get("training_time", 0))

            try:
                plot_buf = Visualizer.create_model_comparison_plot(models)
                if plot_buf:
                    doc.add_heading("Performance Comparison", level=2)
                    doc.add_picture(plot_buf, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"[Chart error: {e}]")
        else:
            doc.add_paragraph("No model results found.")

        # 4. Feature Importance
        doc.add_heading("4. Feature Importance (Top 10)", level=1)
        feat_imp = automl.get("feature_importance", [])
        if feat_imp:
            try:
                plot_buf = Visualizer.create_feature_importance_plot(feat_imp)
                if plot_buf:
                    doc.add_picture(plot_buf, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"[Chart error: {e}]")
        else:
            doc.add_paragraph("Feature importance not available.")

        # 5. Why This Model?
        doc.add_heading("5. Why This Model?", level=1)
        doc.add_paragraph(automl.get("selection_reason", "N/A"))
        for part in automl.get("reason_parts", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(part))

        # 6. Preprocessing Tips
        tips = automl.get("preprocessing_tips", [])
        if tips:
            doc.add_heading("6. Preprocessing Tips", level=1)
            for tip in tips:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(tip))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ── PDF (reportlab) ─────────────────────────────────────────────────────────

    @staticmethod
    def generate_pdf(data: dict) -> io.BytesIO:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
        automl = data.get("automl_results", {})

        # ── Styles ───────────────────────────────────────────────────────────
        base = getSampleStyleSheet()
        BLUE = colors.HexColor("#1E88E5")
        DARK = colors.HexColor("#1a1a2e")

        title_style = ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontSize=20,
            textColor=BLUE,
            spaceAfter=12,
            fontName="Helvetica-Bold",
        )
        h1_style = ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontSize=14,
            textColor=BLUE,
            spaceBefore=14,
            spaceAfter=6,
            fontName="Helvetica-Bold",
        )
        h2_style = ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontSize=12,
            textColor=colors.HexColor("#4FC3F7"),
            spaceBefore=8,
            spaceAfter=4,
            fontName="Helvetica-Bold",
        )
        body = ParagraphStyle(
            "Body",
            parent=base["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=4,
        )
        bullet_style = ParagraphStyle(
            "Bullet",
            parent=body,
            leftIndent=14,
            bulletIndent=6,
        )

        def h(text: str, style=h1_style):
            return Paragraph(text, style)

        def p(text: str):
            return Paragraph(str(text), body)

        def bullet(text: str):
            return Paragraph(f"• {text}", bullet_style)

        def kv(label: str, value) -> Paragraph:
            return Paragraph(f"<b>{label}:</b> {_safe(value)}", body)

        TBL_HEADER = colors.HexColor("#1565C0")
        TBL_ALT = colors.HexColor("#E3F2FD")
        TBL_HIGHLIGHT = colors.HexColor("#C8E6C9")

        def base_table_style(highlight_first=False):
            style = [
                ("BACKGROUND", (0, 0), (-1, 0), TBL_HEADER),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#90CAF9")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, TBL_ALT]),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
            if highlight_first:
                style.append(("BACKGROUND", (0, 1), (-1, 1), TBL_HIGHLIGHT))
            return TableStyle(style)

        # ── Build flowables ───────────────────────────────────────────────────
        story = []
        tmp_files = []  # collect temp files; unlink AFTER doc.build()

        # Title
        story.append(Paragraph(
            f"AutoML Report: {data.get('dataset_name', 'Dataset')}",
            title_style,
        ))
        story.append(Spacer(1, 8))

        # 1. Dataset Overview
        story.append(h("1. Dataset Overview"))
        for label, val in [
            ("Rows", data.get("rows")),
            ("Columns", data.get("columns")),
            ("Task Type", data.get("task_type", "N/A")),
            ("Target Column", data.get("target_column", "N/A")),
        ]:
            story.append(kv(label, val))
        story.append(Spacer(1, 8))

        # 2. Meta-Insights
        story.append(h("2. Meta-Insights"))
        meta = data.get("meta_features", {})
        if meta:
            meta_rows = [
                [Paragraph("<b>Metric</b>", body), Paragraph("<b>Value</b>", body)]
            ]
            for name, val in [
                ("Dimensionality Ratio", meta.get("dimensionality_ratio")),
                ("Class Imbalance Ratio", meta.get("class_imbalance_ratio")),
                ("Avg Feature Correlation", meta.get("avg_feature_correlation")),
                ("Skewness", meta.get("skewness")),
            ]:
                meta_rows.append([p(name), p(_safe(val))])
            meta_tbl = Table(meta_rows, colWidths=[9 * cm, 6 * cm])
            meta_tbl.setStyle(base_table_style())
            story.append(meta_tbl)
        else:
            story.append(p("No meta-features available."))
        story.append(Spacer(1, 8))

        # 3. Algorithm Leaderboard
        story.append(h("3. Algorithm Leaderboard"))
        models = automl.get("algorithms", [])
        if models:
            headers = ["#", "Algorithm", "Accuracy", "F1 Score", "Time (s)"]
            rows = [[Paragraph(f"<b>{h_}</b>", body) for h_ in headers]]
            for i, m in enumerate(models, 1):
                rows.append([
                    p(str(i)),
                    p(m.get("name", "Unknown")),
                    p(f"{m.get('accuracy', 0)}%"),
                    p(f"{m.get('f1_score', 'N/A')}%"),
                    p(_safe(m.get("training_time", 0))),
                ])
            tbl = Table(rows, colWidths=[1 * cm, 5.5 * cm, 3 * cm, 3 * cm, 3 * cm])
            tbl.setStyle(base_table_style(highlight_first=True))
            story.append(tbl)
            story.append(Spacer(1, 8))

            # Model comparison chart
            try:
                plot_buf = Visualizer.create_model_comparison_plot(models)
                if plot_buf:
                    tmp = _chart_to_tmp(plot_buf)
                    if tmp:
                        tmp_files.append(tmp)
                        story.append(h("Performance Comparison", h2_style))
                        story.append(RLImage(tmp, width=15 * cm, height=8 * cm))
                        story.append(Spacer(1, 6))
            except Exception:
                pass
        else:
            story.append(p("No model results found."))
        story.append(Spacer(1, 8))

        # 4. Feature Importance
        story.append(h("4. Feature Importance (Top 10)"))
        feat_imp = automl.get("feature_importance", [])
        if feat_imp:
            try:
                plot_buf = Visualizer.create_feature_importance_plot(feat_imp)
                if plot_buf:
                    tmp = _chart_to_tmp(plot_buf)
                    if tmp:
                        tmp_files.append(tmp)
                        story.append(RLImage(tmp, width=15 * cm, height=8 * cm))
                        story.append(Spacer(1, 6))
            except Exception:
                pass
        else:
            story.append(p("Feature importance not available."))
        story.append(Spacer(1, 8))

        # 5. Why This Model?
        story.append(h("5. Why This Model?"))
        story.append(p(automl.get("selection_reason", "N/A")))
        for part in automl.get("reason_parts", []):
            story.append(bullet(str(part)))
        story.append(Spacer(1, 8))

        # 6. Preprocessing Tips
        tips = automl.get("preprocessing_tips", [])
        if tips:
            story.append(h("6. Preprocessing Tips"))
            for tip in tips:
                story.append(bullet(str(tip)))

        doc.build(story)

        # Clean up temp chart files now that build is complete
        for tmp_path in tmp_files:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

        buf.seek(0)
        return buf
