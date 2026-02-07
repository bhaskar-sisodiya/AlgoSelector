import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# ================================
# 🧠 WORD REPORT GENERATOR
# ================================
def generate_stylish_word_report(dataset_summary, col_info, preprocessing_summary, top_algo, simple_exp, recommendations, class_plot_path):
    word_file = "Stylish_AutoML_Report.docx"
    doc = Document()

    # Title
    title = doc.add_heading("AutoML Report", level=0)
    run = title.runs[0]
    run.font.color.rgb = RGBColor(30, 136, 229)
    run.font.size = Pt(24)

    # Dataset Overview
    doc.add_heading("📘 Dataset Overview", level=1)
    for key, value in dataset_summary.items():
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(value))

    # Columns
    doc.add_heading("📊 Columns Summary", level=1)
    for _, row in col_info.iterrows():
        doc.add_paragraph(f"{row['Column Name']} (Type: {row['Data Type']}, Unique: {row['Unique Values']})")

    # Preprocessing
    if preprocessing_summary:
        doc.add_heading("⚙️ Preprocessing Summary", level=1)
        for key, value in preprocessing_summary.items():
            doc.add_paragraph(f"{key}: {value}")

    # Algorithm Recommendation
    doc.add_heading("🤖 Algorithm Recommendation", level=1)
    doc.add_paragraph(f"Top Algorithm: {top_algo}")
    doc.add_paragraph(f"Justification: {simple_exp}")

    # Alternatives
    if recommendations:
        doc.add_heading("🧩 Alternative Algorithms", level=1)
        for rec in recommendations:
            name = rec.get("name", "Unknown Algorithm")
            best_for = rec.get("best_for", "Not specified")
            doc.add_paragraph(f"• {name} — {best_for}")

    # Chart
    if class_plot_path and os.path.exists(class_plot_path):
        doc.add_heading("📈 Class Distribution", level=1)
        doc.add_picture(class_plot_path, width=Inches(4))

    doc.save(word_file)
    return word_file


# ================================
# 🎨 PDF REPORT GENERATOR
# ================================
def generate_stylish_pdf_report(dataset_summary, col_info, preprocessing_summary, top_algo, simple_exp, recommendations, class_plot_path):
    pdf_file = "Stylish_AutoML_Report.pdf"
    doc = SimpleDocTemplate(pdf_file, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=60, bottomMargin=40)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleStyle", fontSize=20, leading=24, spaceAfter=12,
                              alignment=1, textColor=colors.HexColor("#1E88E5"), fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="SubTitleStyle", fontSize=14, leading=18, spaceAfter=10,
                              textColor=colors.HexColor("#1565C0"), fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="BodyStyle", fontSize=11, leading=16, spaceAfter=8))
    styles.add(ParagraphStyle(name="BoxText", fontSize=11, leading=16, textColor=colors.black))

    flowables = []

    # Title
    flowables.append(Paragraph("AutoML Report", styles["TitleStyle"]))
    flowables.append(Spacer(1, 12))

    # Dataset Overview
    flowables.append(Paragraph("📘 Dataset Overview", styles["SubTitleStyle"]))
    for key, value in dataset_summary.items():
        flowables.append(Paragraph(f"<b>{key}:</b> {value}", styles["BodyStyle"]))
    flowables.append(Spacer(1, 10))

    # Columns
    flowables.append(Paragraph("📊 Columns Summary", styles["SubTitleStyle"]))
    for _, row in col_info.iterrows():
        flowables.append(Paragraph(
            f"{row['Column Name']} (<b>Type:</b> {row['Data Type']}, <b>Unique:</b> {row['Unique Values']})",
            styles["BodyStyle"]
        ))
    flowables.append(Spacer(1, 10))

    # Preprocessing
    if preprocessing_summary:
        flowables.append(Paragraph("⚙️ Preprocessing Summary", styles["SubTitleStyle"]))
        for key, value in preprocessing_summary.items():
            flowables.append(Paragraph(f"<b>{key}:</b> {value}", styles["BodyStyle"]))
        flowables.append(Spacer(1, 10))

    # Algorithm Recommendation
    flowables.append(Paragraph("🤖 Algorithm Recommendation", styles["SubTitleStyle"]))
    flowables.append(Paragraph(f"<b>Top Algorithm:</b> {top_algo}", styles["BodyStyle"]))
    flowables.append(Paragraph(f"<b>Justification:</b> {simple_exp}", styles["BodyStyle"]))
    flowables.append(Spacer(1, 10))

    # Alternatives
    if recommendations:
        flowables.append(Paragraph("🧩 Alternative Algorithms", styles["SubTitleStyle"]))
        alt_text = ""
        for rec in recommendations:
            name = rec.get("name", "Unknown Algorithm")
            best_for = rec.get("best_for", "Not specified")
            alt_text += f"<b>{name}</b> — {best_for}<br/><br/>"
        alt_table = Table([[Paragraph(alt_text, styles["BoxText"])]], colWidths=[400])
        alt_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F3F4F6")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#90A4AE")),
            ("INNERPADDING", (0, 0), (-1, -1), 10)
        ]))
        flowables.append(alt_table)
        flowables.append(Spacer(1, 10))

    # Chart
    if class_plot_path and os.path.exists(class_plot_path):
        flowables.append(Paragraph("📈 Class Distribution", styles["SubTitleStyle"]))
        flowables.append(Image(class_plot_path, width=200, height=200))

    doc.build(flowables)
    return pdf_file
