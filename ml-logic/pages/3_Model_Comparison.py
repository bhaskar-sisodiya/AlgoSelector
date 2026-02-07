# pages/3_Model_Comparison.py

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
import os

st.title("3. Model Comparison & Report Generation")

# Check if dataset is available
if "df" not in st.session_state:
    st.warning("Please upload and process a dataset first on the homepage.")
else:
    df = st.session_state.df
    target_col = st.session_state.target_column if "target_column" in st.session_state else None

    # Dataset Summary
    dataset_summary = {
        "Rows": df.shape[0],
        "Columns": df.shape[1],
        "Missing Values": int(df.isnull().sum().sum())
    }

    st.subheader("📊 Dataset Summary")
    for key, value in dataset_summary.items():
        st.write(f"**{key}:** {value}")

    # Column Information
    col_info = pd.DataFrame({
        "Column Name": df.columns,
        "Data Type": df.dtypes.astype(str),
        "Unique Values": [df[col].nunique() for col in df.columns]
    })

    st.subheader("📑 Column Details")
    st.dataframe(col_info)

    # Recommended Algorithm
    if "recommendation" in st.session_state:
        rec_data = st.session_state.recommendation

        st.subheader("🎯 Recommended Algorithm")

        # Class Distribution Pie Chart (only for categorical target)
        class_plot_path = None
        if target_col and df[target_col].dtype in ['object', 'category', 'bool', 'int64'] and df[target_col].nunique() < 20:
            class_plot_path = "class_distribution.png"
            plt.figure(figsize=(4, 4))
            df[target_col].value_counts().plot.pie(
                autopct='%1.1f%%', colors=sns.color_palette('pastel')
            )
            plt.title(f"Class Distribution: {target_col}")
            plt.tight_layout()
            plt.savefig(class_plot_path)
            plt.close()
            st.image(class_plot_path, caption="Class Distribution")

        # --- Word Report ---
        if st.button("📄 Generate Word Report"):
            word_file = "AutoML_Report.docx"
            doc = Document()
            doc.add_heading("AutoML Report", 0)

            doc.add_heading("Dataset Overview", level=1)
            for key, value in dataset_summary.items():
                doc.add_paragraph(f"{key}: {value}")

            doc.add_heading("Columns", level=1)
            for _, row in col_info.iterrows():
                doc.add_paragraph(f"{row['Column Name']} (Type: {row['Data Type']}, Unique: {row['Unique Values']})")

            doc.add_heading("Algorithm Recommendation", level=1)
            doc.add_paragraph(f"Algorithm: {rec['algorithm']}")
            doc.add_paragraph(f"Justification: {rec['simple_explanation']}")

            if class_plot_path and os.path.exists(class_plot_path):
                doc.add_heading("Graphs", level=1)
                doc.add_picture(class_plot_path, width=Inches(4))

            doc.save(word_file)
            st.success(f"✅ Word report generated: {word_file}")
            with open(word_file, "rb") as f:
                st.download_button(
                    label="⬇️ Download Word Report",
                    data=f,
                    file_name=word_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        # --- PDF Report ---
        if st.button("📑 Generate PDF Report"):
            pdf_file = "AutoML_Report.pdf"
            doc_pdf = SimpleDocTemplate(pdf_file, pagesize=letter)
            styles = getSampleStyleSheet()
            flowables = []

            flowables.append(Paragraph("AutoML Report", styles['Title']))
            flowables.append(Spacer(1, 12))

            flowables.append(Paragraph("Dataset Overview", styles['Heading2']))
            for key, value in dataset_summary.items():
                flowables.append(Paragraph(f"{key}: {value}", styles['Normal']))
            flowables.append(Spacer(1, 12))

            flowables.append(Paragraph("Columns", styles['Heading2']))
            for _, row in col_info.iterrows():
                flowables.append(Paragraph(f"{row['Column Name']} (Type: {row['Data Type']}, Unique: {row['Unique Values']})", styles['Normal']))
            flowables.append(Spacer(1, 12))

            flowables.append(Paragraph("Algorithm Recommendation", styles['Heading2']))
            flowables.append(Paragraph(f"Algorithm: {rec['algorithm']}", styles['Normal']))
            flowables.append(Paragraph(f"Justification: {rec['simple_explanation']}", styles['Normal']))
            flowables.append(Spacer(1, 12))

            # Include class distribution chart if exists
            if class_plot_path and os.path.exists(class_plot_path):
                flowables.append(Paragraph("Class Distribution", styles['Heading2']))
                flowables.append(Image(class_plot_path, width=200, height=200))

            doc_pdf.build(flowables)
            st.success(f"✅ PDF report generated: {pdf_file}")
            with open(pdf_file, "rb") as f:
                st.download_button(
                    label="⬇️ Download PDF Report",
                    data=f,
                    file_name=pdf_file,
                    mime="application/pdf"
                )

    else:
        st.info("Please go to '2. Algorithm Recommendation' page to get the recommendation.")
